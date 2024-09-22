# print ("Welcome to New Project!")
# print("Here is the workspace for the project")

# import tkinter as tk
# from tkinter import ttk, messagebox, simpledialog

# class MedicalReportSystem:
#     def __init__(self, master):
#         self.master = master
#         self.master.title("Medical Report System")
#         self.master.geometry("600x400")

#         self.reports = []

#         # Create and set up the notebook
#         self.notebook = ttk.Notebook(self.master)
#         self.notebook.pack(expand=True, fill="both", padx=10, pady=10)

#         # Create tabs
#         self.create_add_report_tab()
#         self.create_view_reports_tab()

#     def create_add_report_tab(self):
#         add_frame = ttk.Frame(self.notebook)
#         self.notebook.add(add_frame, text="Add Report")

#         # Patient Name
#         ttk.Label(add_frame, text="Patient Name:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
#         self.patient_name = ttk.Entry(add_frame, width=40)
#         self.patient_name.grid(row=0, column=1, padx=5, pady=5)

#         # Date
#         ttk.Label(add_frame, text="Date:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
#         self.date = ttk.Entry(add_frame, width=40)
#         self.date.grid(row=1, column=1, padx=5, pady=5)

#         # Diagnosis
#         ttk.Label(add_frame, text="Diagnosis:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
#         self.diagnosis = tk.Text(add_frame, width=40, height=5)
#         self.diagnosis.grid(row=2, column=1, padx=5, pady=5)

#         # Treatment
#         ttk.Label(add_frame, text="Treatment:").grid(row=3, column=0, padx=5, pady=5, sticky="w")
#         self.treatment = tk.Text(add_frame, width=40, height=5)
#         self.treatment.grid(row=3, column=1, padx=5, pady=5)

#         # Submit Button
#         submit_btn = ttk.Button(add_frame, text="Submit Report", command=self.submit_report)
#         submit_btn.grid(row=4, column=1, pady=20)

#     def create_view_reports_tab(self):
#         view_frame = ttk.Frame(self.notebook)
#         self.notebook.add(view_frame, text="View Reports")

#         self.report_listbox = tk.Listbox(view_frame, width=60)
#         self.report_listbox.pack(padx=10, pady=10, fill="both", expand=True)

#         view_btn = ttk.Button(view_frame, text="View Selected Report", command=self.view_report)
#         view_btn.pack(pady=10)

#     def submit_report(self):
#         name = self.patient_name.get()
#         date = self.date.get()
#         diagnosis = self.diagnosis.get("1.0", "end-1c")
#         treatment = self.treatment.get("1.0", "end-1c")

#         if name and date and diagnosis and treatment:
#             report = f"{name} - {date}"
#             self.reports.append({
#                 "name": name,
#                 "date": date,
#                 "diagnosis": diagnosis,
#                 "treatment": treatment
#             })
#             self.report_listbox.insert(tk.END, report)
#             messagebox.showinfo("Success", "Report submitted successfully!")
#             self.clear_fields()
#         else:
#             messagebox.showerror("Error", "All fields are required!")

#     def view_report(self):
#         selection = self.report_listbox.curselection()
#         if selection:
#             index = selection[0]
#             report = self.reports[index]
#             report_text = f"Patient: {report['name']}\n\n"
#             report_text += f"Date: {report['date']}\n\n"
#             report_text += f"Diagnosis:\n{report['diagnosis']}\n\n"
#             report_text += f"Treatment:\n{report['treatment']}"
            
#             report_window = tk.Toplevel(self.master)
#             report_window.title(f"Report: {report['name']} - {report['date']}")
#             report_window.geometry("400x300")
            
#             text_widget = tk.Text(report_window, wrap=tk.WORD, padx=10, pady=10)
#             text_widget.insert(tk.END, report_text)
#             text_widget.config(state=tk.DISABLED)
#             text_widget.pack(expand=True, fill="both")
#         else:
#             messagebox.showerror("Error", "Please select a report to view.")

#     def clear_fields(self):
#         self.patient_name.delete(0, tk.END)
#         self.date.delete(0, tk.END)
#         self.diagnosis.delete("1.0", tk.END)
#         self.treatment.delete("1.0", tk.END)

# if __name__ == "__main__":
#     root = tk.Tk()
#     app = MedicalReportSystem(root)
#     root.mainloop()
 
import docx 
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)


records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')